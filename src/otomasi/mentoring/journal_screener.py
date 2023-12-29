import os
from docx import Document
import hashlib
import re
from collections import defaultdict
from tqdm import tqdm
import concurrent.futures
import pandas as pd


class ContentAnalyzer:
    """Handles content analysis and validation"""

    TEMPLATE_HEADERS_EN = [
        "Adoration",
        "Confession",
        "Thanksgiving",
        "Supplication",
        "Prayer Points",
        "Answer:",
        "Reflection",
        "Application",
    ]

    TEMPLATE_HEADERS_ID = [
        "Pokok Doa",
        "Pengakuan",
        "Pengucapan Syukur",
        "Permohonan",
        "Jawab:",
        "Refleksi",
        "Aplikasi",
    ]

    def __init__(self):
        self.template_headers = set(self.TEMPLATE_HEADERS_EN + self.TEMPLATE_HEADERS_ID)
        self._check_nltk_installation()

    def _check_nltk_installation(self):
        """Verify NLTK is properly installed"""
        try:
            import nltk

            nltk.download("punkt_tab")

            from nltk.tokenize import word_tokenize

            word_tokenize("Test sentence")
        except (LookupError, ImportError) as e:
            print("NLTK data not found. Please run setup_nltk.py first.")
            print("Error:", str(e))
            raise SystemExit(1)

    def count_meaningful_words(self, text):
        """Count words excluding template headers and common words"""
        # Remove template headers
        for header in self.template_headers:
            text = text.replace(header, "")

        # Simple word count after cleaning
        words = [w for w in text.split() if len(w) > 1]
        return len(words)

    def detect_repetitive_patterns(self, text):
        """Detect copy-pasted content within the document"""
        # Split into paragraphs
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

        # Compare paragraphs for similarity
        suspicious_patterns = []
        for i in range(len(paragraphs)):
            for j in range(i + 1, len(paragraphs)):
                # Basic similarity check - could be improved
                if len(paragraphs[i]) > 50 and paragraphs[i] == paragraphs[j]:
                    suspicious_patterns.append(
                        {
                            "paragraph1": paragraphs[i][:100] + "...",
                            "paragraph2": paragraphs[j][:100] + "...",
                            "similarity": 1.0,
                        }
                    )

        return suspicious_patterns


class FileProcessor:
    """Handles file operations and checking"""

    def __init__(self):
        self.content_analyzer = ContentAnalyzer()
        self.student_submissions = defaultdict(list)

    def extract_file_info(self, filename):
        """Extract date and student info from filename"""
        try:
            # More flexible pattern matching for file naming
            date_pattern = r"(\d{1,2}\s*[A-Za-z]{3,4})"  # Matches dates like "11 Des", "1 Dec", etc.
            id_pattern = r"(\d{7,8})"  # Matches student IDs of 7-8 digits
            name_pattern = (
                r"([A-Za-z\s\-\.]+)"  # Matches names with spaces, hyphens, and dots
            )

            # First try to find student ID
            id_match = re.search(id_pattern, filename)
            if not id_match:
                return None

            # Then look for date and name
            date_match = re.search(date_pattern, filename)
            name_match = re.search(name_pattern, filename)

            if not (date_match and name_match):
                return None

            # Clean up extracted data
            student_id = id_match.group(1)
            date = date_match.group(1).strip()
            name = name_match.group(1).strip()

            # Validate extracted data
            if len(student_id) < 7 or not date or len(name) < 2:
                return None

            return {
                "date": date,
                "student_id": student_id,
                "name": name,
                "filename": filename,
            }
        except Exception as e:
            print(f"Error processing filename '{filename}': {str(e)}")
            return None

    def check_file_integrity(self, file_path):
        """Verify file is a valid DOCX"""
        try:
            doc = Document(file_path)
            # Try to read some content to verify
            _ = [p.text for p in doc.paragraphs[:5]]
            return True, None
        except Exception as e:
            return False, str(e)

    def get_file_content(self, file_path):
        """Extract text content from DOCX file"""
        try:
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            return None

    def process_file(self, file_path):
        """Process a single file and return analysis results"""
        filename = os.path.basename(file_path)
        file_info = self.extract_file_info(filename)

        if not file_info:
            return {"filename": filename, "issues": ["Invalid filename format"]}

        # Check file integrity
        is_valid, error = self.check_file_integrity(file_path)
        if not is_valid:
            return {
                "filename": filename,
                "student_id": file_info["student_id"],
                "name": file_info["name"],
                "issues": [f"File integrity error: {error}"],
            }

        # Get content
        content = self.get_file_content(file_path)
        if not content:
            return {
                "filename": filename,
                "student_id": file_info["student_id"],
                "name": file_info["name"],
                "issues": ["Unable to read file content"],
            }

        # Analyze content
        word_count = self.content_analyzer.count_meaningful_words(content)
        repetitive_patterns = self.content_analyzer.detect_repetitive_patterns(content)

        # Track for duplicate detection
        content_hash = hashlib.md5(content.encode()).hexdigest()
        self.student_submissions[file_info["student_id"]].append(
            {"date": file_info["date"], "hash": content_hash, "filename": filename}
        )

        # Collect issues
        issues = []
        if word_count < 70:
            issues.append(f"Low word count ({word_count} words)")
        if repetitive_patterns:
            issues.append(f"Found {len(repetitive_patterns)} repetitive patterns")

        return {
            "filename": filename,
            "student_id": file_info["student_id"],
            "name": file_info["name"],
            "date": file_info["date"],
            "word_count": word_count,
            "content_hash": content_hash,
            "repetitive_patterns": repetitive_patterns,
            "issues": issues,
        }


class JournalValidator:
    """Main class for journal validation"""

    def __init__(self):
        self.file_processor = FileProcessor()

    def find_duplicate_submissions(self):
        """Find duplicate submissions across different dates for each student"""
        duplicates = {}  # Use dict to store first duplicate instance per student

        for student_id, submissions in self.file_processor.student_submissions.items():
            if len(submissions) < 2:
                continue

            # Sort submissions by date
            sorted_submissions = sorted(submissions, key=lambda x: x["date"])

            # Find first instance of duplicate for this student
            for i in range(len(sorted_submissions) - 1):
                if sorted_submissions[i]["hash"] == sorted_submissions[i + 1]["hash"]:
                    if student_id not in duplicates:  # Only store first instance
                        duplicates[student_id] = {
                            "student_id": student_id,
                            "file1": sorted_submissions[i]["filename"],
                            "date1": sorted_submissions[i]["date"],
                            "file2": sorted_submissions[i + 1]["filename"],
                            "date2": sorted_submissions[i + 1]["date"],
                        }
                    break  # Stop after finding first duplicate for this student

        return list(duplicates.values())

    def process_directory(self, directory_path):
        """Process all DOCX files in directory"""
        files = [f for f in os.listdir(directory_path) if f.endswith(".docx")]
        results = []
        invalid_files = []

        print(f"\nProcessing {len(files)} files...")

        # Process files in parallel
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            futures = [
                executor.submit(
                    self.file_processor.process_file, os.path.join(directory_path, file)
                )
                for file in files
            ]

            for future in tqdm(
                concurrent.futures.as_completed(futures), total=len(files)
            ):
                try:
                    result = future.result()
                    if "student_id" not in result:
                        invalid_files.append(result["filename"])
                        continue
                    results.append(result)
                except Exception as e:
                    print(f"Error processing file: {str(e)}")

        # Find duplicates after processing all files
        duplicates = self.find_duplicate_submissions()

        # Add duplicate issues to results
        for dup in duplicates:
            for result in results:
                if result.get("student_id") == dup["student_id"]:
                    result["issues"].append(
                        f"Duplicate submission with dates: {dup['date1']} and {dup['date2']}"
                    )

        # Add invalid files to results summary
        if invalid_files:
            print(f"\nFound {len(invalid_files)} files with invalid format:")
            for file in invalid_files[:10]:  # Show first 10 invalid files
                print(f"- {file}")
            if len(invalid_files) > 10:
                print(f"... and {len(invalid_files) - 10} more")

        return results

    def save_results(self, results, output_file="journal_validation_results.xlsx"):
        """Save results to Excel file"""
        # Track issues per student to avoid duplicates
        student_issues = defaultdict(dict)  # student_id -> issue_type -> issue_details
        patterns_data = []

        # Process all results
        for result in results:
            student_id = result.get("student_id", "Unknown")
            name = result.get("name", "Unknown")

            # Process each issue
            for issue in result.get("issues", []):
                # Determine issue type
                issue_type = None
                if "Duplicate submission" in issue:
                    issue_type = "duplicate"
                elif "Low word count" in issue:
                    issue_type = "low_word_count"
                elif "repetitive patterns" in issue:
                    issue_type = "repetitive"
                else:
                    issue_type = "other"

                # Only store if this type of issue hasn't been recorded for this student
                if issue_type not in student_issues[student_id]:
                    student_issues[student_id][issue_type] = {
                        "Student ID": student_id,
                        "Name": name,
                        "Filename": result["filename"],
                        "Issue": issue,
                        "Word Count": result.get("word_count", 0),
                    }

            # Process repetitive patterns
            if (
                result.get("repetitive_patterns")
                and student_id not in student_issues["repetitive"]
            ):
                patterns_data.append(
                    {
                        "Student ID": student_id,
                        "Name": name,
                        "Filename": result["filename"],
                        "Similar Text 1": result["repetitive_patterns"][0][
                            "paragraph1"
                        ],
                        "Similar Text 2": result["repetitive_patterns"][0][
                            "paragraph2"
                        ],
                        "Similarity Score": result["repetitive_patterns"][0][
                            "similarity"
                        ],
                    }
                )

        # Prepare final issues data
        issues_data = []
        for student_issues_dict in student_issues.values():
            issues_data.extend(student_issues_dict.values())

        # Create Excel writer
        with pd.ExcelWriter(output_file, engine="openpyxl") as writer:
            # Write issues sheet
            pd.DataFrame(issues_data).to_excel(writer, sheet_name="Issues", index=False)

            # Write patterns sheet if any found
            if patterns_data:
                pd.DataFrame(patterns_data).to_excel(
                    writer, sheet_name="Repetitive Patterns", index=False
                )


def main(journal_dir: str, output_file: str):
    validator = JournalValidator()
    # Process all files
    print("Starting journal validation...")
    results = validator.process_directory(journal_dir)

    # Save results
    validator.save_results(results, output_file)

    print(f"\nValidation complete! Results saved to {output_file}")
    print("Please check the Excel file for:")
    print(
        "- 'Issues' sheet: All detected issues including duplicates and low word count"
    )
    if any(r.get("repetitive_patterns") for r in results):
        print(
            "- 'Repetitive Patterns' sheet: Detected copy-paste patterns within documents"
        )
